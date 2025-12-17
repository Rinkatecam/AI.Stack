"""
title: Knowledge Base Tool
version: 1.0.0
description: Experience database with auto-indexing of DOCX files and image comparison capabilities.
author: AI.STACK
author_url: https://github.com/Rinkatecam/aistack
requirements: pydantic, python-docx, pillow, qdrant-client

# SYSTEM PROMPT FOR AI
# ====================
# Organizational knowledge base with document indexing and image comparison.
#
# CAPABILITIES:
#   - Auto-index DOCX files from configured directories
#   - Semantic search across indexed documents
#   - Add/retrieve experiences and solutions
#   - Image comparison for QS inspections
#
# WORKFLOW:
#   1. index_folder(path) - Index documents
#   2. search_knowledge(query) - Find relevant info
#   3. add_experience(title, content, tags) - Add new entry
#   4. compare_images(img1, img2) - Visual comparison
"""

import os
import json
import hashlib
from datetime import datetime
from typing import Optional, List, Dict, Any
from pydantic import BaseModel, Field

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from PIL import Image
    import io
    PILLOW_AVAILABLE = True
except ImportError:
    PILLOW_AVAILABLE = False

try:
    from qdrant_client import QdrantClient
    from qdrant_client.models import Distance, VectorParams, PointStruct
    QDRANT_AVAILABLE = True
except ImportError:
    QDRANT_AVAILABLE = False

try:
    import requests
    REQUESTS_AVAILABLE = True
except ImportError:
    REQUESTS_AVAILABLE = False


class Tools:
    class Valves(BaseModel):
        knowledge_dir: str = Field(
            default="/data/knowledge",
            description="Base directory for knowledge base files"
        )
        auto_index_dirs: str = Field(
            default="/data/knowledge/documents",
            description="Comma-separated directories to auto-index"
        )
        qdrant_host: str = Field(
            default="aistack-vector",
            description="Qdrant vector database host"
        )
        qdrant_port: int = Field(
            default=6333,
            description="Qdrant port"
        )
        embedding_model: str = Field(
            default="nomic-embed-text",
            description="Ollama embedding model"
        )
        ollama_host: str = Field(
            default="ollama",
            description="Ollama host for embeddings"
        )
        collection_name: str = Field(
            default="knowledge_base",
            description="Qdrant collection name"
        )
        embedding_dim: int = Field(
            default=768,
            description="Embedding vector dimension"
        )
        chunk_size: int = Field(
            default=500,
            description="Text chunk size for indexing"
        )
        image_similarity_threshold: float = Field(
            default=0.85,
            description="Threshold for image similarity (0-1)"
        )

    def __init__(self):
        self.citation = False
        self.valves = self.Valves()
        self._ensure_directories()
        self._experiences_file = os.path.join(self.valves.knowledge_dir, "experiences.json")

    def _ensure_directories(self):
        """Create necessary directories."""
        os.makedirs(self.valves.knowledge_dir, exist_ok=True)
        for dir_path in self.valves.auto_index_dirs.split(','):
            dir_path = dir_path.strip()
            if dir_path:
                os.makedirs(dir_path, exist_ok=True)

    def _get_qdrant_client(self) -> Optional[Any]:
        """Get Qdrant client."""
        if not QDRANT_AVAILABLE:
            return None
        try:
            return QdrantClient(host=self.valves.qdrant_host, port=self.valves.qdrant_port)
        except:
            return None

    def _get_embedding(self, text: str) -> Optional[List[float]]:
        """Get embedding from Ollama."""
        if not REQUESTS_AVAILABLE:
            return None
        try:
            url = f"http://{self.valves.ollama_host}:11434/api/embeddings"
            response = requests.post(url, json={
                "model": self.valves.embedding_model,
                "prompt": text
            }, timeout=30)
            if response.status_code == 200:
                return response.json().get("embedding")
        except:
            pass
        return None

    def _ensure_collection(self, client):
        """Ensure Qdrant collection exists."""
        collections = [c.name for c in client.get_collections().collections]
        if self.valves.collection_name not in collections:
            client.create_collection(
                collection_name=self.valves.collection_name,
                vectors_config=VectorParams(
                    size=self.valves.embedding_dim,
                    distance=Distance.COSINE
                )
            )

    def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX file."""
        if not DOCX_AVAILABLE:
            return ""
        try:
            doc = DocxDocument(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text.strip())

            return "\n".join(paragraphs)
        except:
            return ""

    def _chunk_text(self, text: str, chunk_size: int = None) -> List[str]:
        """Split text into chunks."""
        chunk_size = chunk_size or self.valves.chunk_size
        words = text.split()
        chunks = []
        current_chunk = []
        current_size = 0

        for word in words:
            current_chunk.append(word)
            current_size += len(word) + 1

            if current_size >= chunk_size:
                chunks.append(' '.join(current_chunk))
                current_chunk = []
                current_size = 0

        if current_chunk:
            chunks.append(' '.join(current_chunk))

        return chunks

    def _load_experiences(self) -> List[Dict]:
        """Load experiences from JSON file."""
        if os.path.exists(self._experiences_file):
            try:
                with open(self._experiences_file, 'r', encoding='utf-8') as f:
                    return json.load(f)
            except:
                pass
        return []

    def _save_experiences(self, experiences: List[Dict]):
        """Save experiences to JSON file."""
        with open(self._experiences_file, 'w', encoding='utf-8') as f:
            json.dump(experiences, f, indent=2, ensure_ascii=False)

    def index_folder(self, path: str = "", file_types: str = "docx") -> str:
        """
        Index documents from a folder into the knowledge base.

        Args:
            path: Folder path to index (default: auto_index_dirs from settings)
            file_types: File types to index (default: docx)

        Returns:
            Indexing results.
        """
        if not path:
            path = self.valves.auto_index_dirs.split(',')[0].strip()

        if not os.path.exists(path):
            return f"Error: Path not found: {path}"

        if not QDRANT_AVAILABLE:
            return "Error: qdrant-client not installed. Run: pip install qdrant-client"

        client = self._get_qdrant_client()
        if not client:
            return "Error: Could not connect to Qdrant. Check configuration."

        self._ensure_collection(client)

        indexed = 0
        errors = 0
        file_list = []

        extensions = [f".{ext.strip()}" for ext in file_types.split(',')]

        for root, dirs, files in os.walk(path):
            for file in files:
                if any(file.lower().endswith(ext) for ext in extensions):
                    file_path = os.path.join(root, file)
                    file_list.append(file_path)

        result = [
            f"**Indexing Knowledge Base**",
            f"Path: {path}",
            f"Files found: {len(file_list)}",
            "=" * 50,
            ""
        ]

        for file_path in file_list:
            try:
                # Extract text based on file type
                if file_path.lower().endswith('.docx'):
                    text = self._extract_docx_text(file_path)
                else:
                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        text = f.read()

                if not text.strip():
                    result.append(f"⚠️ {os.path.basename(file_path)}: Empty or unreadable")
                    continue

                # Chunk text
                chunks = self._chunk_text(text)

                # Index each chunk
                for i, chunk in enumerate(chunks):
                    embedding = self._get_embedding(chunk)
                    if embedding:
                        doc_id = hashlib.md5(f"{file_path}_{i}".encode()).hexdigest()
                        client.upsert(
                            collection_name=self.valves.collection_name,
                            points=[PointStruct(
                                id=doc_id,
                                vector=embedding,
                                payload={
                                    "source": file_path,
                                    "filename": os.path.basename(file_path),
                                    "chunk": i,
                                    "text": chunk,
                                    "indexed_at": datetime.now().isoformat()
                                }
                            )]
                        )

                indexed += 1
                result.append(f"✅ {os.path.basename(file_path)}: {len(chunks)} chunks")

            except Exception as e:
                errors += 1
                result.append(f"❌ {os.path.basename(file_path)}: {str(e)}")

        result.extend([
            "",
            f"**Summary:** {indexed} files indexed, {errors} errors"
        ])

        return "\n".join(result)

    def search_knowledge(self, query: str, limit: int = 5) -> str:
        """
        Search the knowledge base for relevant information.

        Args:
            query: Search query
            limit: Maximum results to return

        Returns:
            Relevant documents and experiences.
        """
        results_list = []

        # Search vector database
        if QDRANT_AVAILABLE:
            client = self._get_qdrant_client()
            if client:
                embedding = self._get_embedding(query)
                if embedding:
                    try:
                        search_results = client.search(
                            collection_name=self.valves.collection_name,
                            query_vector=embedding,
                            limit=limit
                        )
                        for r in search_results:
                            results_list.append({
                                'type': 'document',
                                'score': r.score,
                                'source': r.payload.get('filename', 'Unknown'),
                                'text': r.payload.get('text', '')[:300],
                                'path': r.payload.get('source', '')
                            })
                    except:
                        pass

        # Search experiences (keyword-based)
        experiences = self._load_experiences()
        query_lower = query.lower()

        for exp in experiences:
            score = 0
            title = exp.get('title', '').lower()
            content = exp.get('content', '').lower()
            tags = [t.lower() for t in exp.get('tags', [])]

            # Simple keyword matching
            for word in query_lower.split():
                if word in title:
                    score += 3
                if word in content:
                    score += 1
                if word in tags:
                    score += 2

            if score > 0:
                results_list.append({
                    'type': 'experience',
                    'score': score / 10,  # Normalize
                    'title': exp.get('title'),
                    'text': exp.get('content', '')[:300],
                    'tags': exp.get('tags', []),
                    'created': exp.get('created_at', '')
                })

        # Sort by score
        results_list.sort(key=lambda x: x['score'], reverse=True)
        results_list = results_list[:limit]

        # Format output
        result = [
            f"**Knowledge Base Search: '{query}'**",
            "=" * 50,
            ""
        ]

        if not results_list:
            result.append("No results found.")
            result.append("\nTry:")
            result.append("  • Different keywords")
            result.append("  • index_folder() to add more documents")
            result.append("  • add_experience() to add solutions")
            return "\n".join(result)

        for i, r in enumerate(results_list, 1):
            if r['type'] == 'document':
                result.append(f"**{i}. 📄 {r['source']}** (score: {r['score']:.2f})")
                result.append(f"   {r['text']}...")
                if r.get('path'):
                    result.append(f"   Path: {r['path']}")
            else:
                result.append(f"**{i}. 💡 {r['title']}** (score: {r['score']:.2f})")
                result.append(f"   {r['text']}...")
                if r.get('tags'):
                    result.append(f"   Tags: {', '.join(r['tags'])}")
            result.append("")

        return "\n".join(result)

    def add_experience(self, title: str, content: str, tags: str = "", category: str = "general") -> str:
        """
        Add a new experience/solution to the knowledge base.

        Args:
            title: Title of the experience
            content: Detailed description/solution
            tags: Comma-separated tags for categorization
            category: Category (e.g., QS, R&D, IT, RA, HR)

        Returns:
            Confirmation message.
        """
        experiences = self._load_experiences()

        # Parse tags
        tag_list = [t.strip() for t in tags.split(',') if t.strip()]

        # Create experience entry
        experience = {
            'id': hashlib.md5(f"{title}{datetime.now().isoformat()}".encode()).hexdigest()[:12],
            'title': title,
            'content': content,
            'tags': tag_list,
            'category': category,
            'created_at': datetime.now().isoformat(),
            'updated_at': datetime.now().isoformat()
        }

        experiences.append(experience)
        self._save_experiences(experiences)

        # Also index in vector DB if available
        if QDRANT_AVAILABLE:
            client = self._get_qdrant_client()
            if client:
                embedding = self._get_embedding(f"{title} {content}")
                if embedding:
                    try:
                        self._ensure_collection(client)
                        client.upsert(
                            collection_name=self.valves.collection_name,
                            points=[PointStruct(
                                id=experience['id'],
                                vector=embedding,
                                payload={
                                    "type": "experience",
                                    "title": title,
                                    "text": content,
                                    "tags": tag_list,
                                    "category": category,
                                    "indexed_at": datetime.now().isoformat()
                                }
                            )]
                        )
                    except:
                        pass

        result = [
            "✅ Experience added to knowledge base!",
            "",
            f"**Title:** {title}",
            f"**Category:** {category}",
            f"**Tags:** {', '.join(tag_list) if tag_list else 'None'}",
            f"**ID:** {experience['id']}",
            "",
            f"Total experiences: {len(experiences)}"
        ]

        return "\n".join(result)

    def list_experiences(self, category: str = "", limit: int = 20) -> str:
        """
        List experiences in the knowledge base.

        Args:
            category: Filter by category (optional)
            limit: Maximum to show

        Returns:
            List of experiences.
        """
        experiences = self._load_experiences()

        if category:
            experiences = [e for e in experiences if e.get('category', '').lower() == category.lower()]

        experiences = experiences[-limit:]  # Most recent

        result = [
            "**Knowledge Base Experiences**",
            "=" * 50,
            f"Total: {len(self._load_experiences())} | Showing: {len(experiences)}",
            ""
        ]

        if not experiences:
            result.append("No experiences found.")
            result.append("\nUse add_experience() to add solutions and learnings.")
            return "\n".join(result)

        for exp in experiences:
            result.append(f"**{exp.get('title', 'Untitled')}**")
            result.append(f"  Category: {exp.get('category', 'N/A')} | ID: {exp.get('id', 'N/A')}")
            if exp.get('tags'):
                result.append(f"  Tags: {', '.join(exp['tags'])}")
            result.append(f"  {exp.get('content', '')[:100]}...")
            result.append("")

        return "\n".join(result)

    def compare_images(self, image1_path: str, image2_path: str) -> str:
        """
        Compare two images for similarity (useful for QS visual inspection).

        Args:
            image1_path: Path to first image (reference)
            image2_path: Path to second image (sample)

        Returns:
            Comparison result with similarity score.
        """
        if not PILLOW_AVAILABLE:
            return "Error: Pillow not installed. Run: pip install pillow"

        if not os.path.exists(image1_path):
            return f"Error: Image not found: {image1_path}"
        if not os.path.exists(image2_path):
            return f"Error: Image not found: {image2_path}"

        try:
            # Load images
            img1 = Image.open(image1_path).convert('RGB')
            img2 = Image.open(image2_path).convert('RGB')

            # Resize to same dimensions for comparison
            target_size = (256, 256)
            img1_resized = img1.resize(target_size, Image.Resampling.LANCZOS)
            img2_resized = img2.resize(target_size, Image.Resampling.LANCZOS)

            # Calculate histogram similarity
            hist1 = img1_resized.histogram()
            hist2 = img2_resized.histogram()

            # Normalize histograms
            sum1 = sum(hist1)
            sum2 = sum(hist2)
            hist1_norm = [x / sum1 for x in hist1]
            hist2_norm = [x / sum2 for x in hist2]

            # Calculate similarity (correlation coefficient)
            mean1 = sum(hist1_norm) / len(hist1_norm)
            mean2 = sum(hist2_norm) / len(hist2_norm)

            numerator = sum((h1 - mean1) * (h2 - mean2) for h1, h2 in zip(hist1_norm, hist2_norm))
            denom1 = sum((h - mean1) ** 2 for h in hist1_norm) ** 0.5
            denom2 = sum((h - mean2) ** 2 for h in hist2_norm) ** 0.5

            if denom1 * denom2 > 0:
                similarity = numerator / (denom1 * denom2)
            else:
                similarity = 0

            # Determine pass/fail based on threshold
            threshold = self.valves.image_similarity_threshold
            passed = similarity >= threshold

            result = [
                "**Image Comparison Result**",
                "=" * 50,
                "",
                f"Reference: {os.path.basename(image1_path)}",
                f"  Size: {img1.size[0]}x{img1.size[1]}",
                "",
                f"Sample: {os.path.basename(image2_path)}",
                f"  Size: {img2.size[0]}x{img2.size[1]}",
                "",
                f"**Similarity Score:** {similarity:.2%}",
                f"**Threshold:** {threshold:.2%}",
                "",
            ]

            if passed:
                result.append("✅ **PASS** - Images are sufficiently similar")
            else:
                result.append("❌ **FAIL** - Images differ significantly")
                result.append("")
                result.append("Possible reasons for difference:")
                result.append("  • Different lighting conditions")
                result.append("  • Position/orientation variance")
                result.append("  • Actual defect or variation")

            return "\n".join(result)

        except Exception as e:
            return f"Error comparing images: {str(e)}"

    def get_statistics(self) -> str:
        """Get knowledge base statistics."""
        experiences = self._load_experiences()

        # Count by category
        categories = {}
        for exp in experiences:
            cat = exp.get('category', 'unknown')
            categories[cat] = categories.get(cat, 0) + 1

        # Count indexed documents
        doc_count = 0
        if QDRANT_AVAILABLE:
            client = self._get_qdrant_client()
            if client:
                try:
                    info = client.get_collection(self.valves.collection_name)
                    doc_count = info.points_count
                except:
                    pass

        result = [
            "**Knowledge Base Statistics**",
            "=" * 50,
            "",
            f"**Experiences:** {len(experiences)}",
            f"**Indexed chunks:** {doc_count}",
            "",
            "**By Category:**",
        ]

        for cat, count in sorted(categories.items()):
            result.append(f"  • {cat}: {count}")

        result.extend([
            "",
            f"Knowledge directory: {self.valves.knowledge_dir}",
            f"Auto-index dirs: {self.valves.auto_index_dirs}",
        ])

        return "\n".join(result)

    def check_availability(self) -> str:
        """Check tool dependencies and configuration."""
        status = ["**Knowledge Base Tool - Status**", "=" * 50, ""]

        status.append(f"{'✅' if DOCX_AVAILABLE else '❌'} python-docx: {'Available' if DOCX_AVAILABLE else 'Not installed'}")
        status.append(f"{'✅' if PILLOW_AVAILABLE else '❌'} Pillow: {'Available' if PILLOW_AVAILABLE else 'Not installed'}")
        status.append(f"{'✅' if QDRANT_AVAILABLE else '❌'} qdrant-client: {'Available' if QDRANT_AVAILABLE else 'Not installed'}")
        status.append(f"{'✅' if REQUESTS_AVAILABLE else '❌'} requests: {'Available' if REQUESTS_AVAILABLE else 'Not installed'}")

        # Check Qdrant connection
        if QDRANT_AVAILABLE:
            client = self._get_qdrant_client()
            if client:
                try:
                    client.get_collections()
                    status.append(f"✅ Qdrant connection: OK ({self.valves.qdrant_host}:{self.valves.qdrant_port})")
                except:
                    status.append(f"❌ Qdrant connection: Failed")
            else:
                status.append(f"❌ Qdrant connection: Could not create client")

        status.extend([
            "",
            "**Configuration:**",
            f"  Knowledge dir: {self.valves.knowledge_dir}",
            f"  Auto-index: {self.valves.auto_index_dirs}",
            f"  Collection: {self.valves.collection_name}",
            "",
            "**Functions:**",
            "  • index_folder(path) - Index documents",
            "  • search_knowledge(query) - Search knowledge base",
            "  • add_experience(title, content, tags) - Add experience",
            "  • compare_images(img1, img2) - Compare images",
            "  • list_experiences() - List all experiences",
        ])

        return "\n".join(status)
