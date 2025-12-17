"""
title: Document Templates Tool
version: 1.0.0
description: Process DOCX templates with {{ PLACEHOLDER }} syntax. AI fills placeholders with user confirmation.
author: AI.STACK
author_url: https://github.com/Rinkatecam/aistack
requirements: pydantic, python-docx

# SYSTEM PROMPT FOR AI
# ====================
# Process Word document templates with {{ PLACEHOLDER }} syntax.
# When filling templates, ASK FOR CONFIRMATION if unsure what to fill.
#
# WORKFLOW:
#   1. list_templates() - See available templates
#   2. get_placeholders("template.docx") - See what needs to be filled
#   3. fill_template("template.docx", {"placeholder": "value"}) - Fill and save
#
# IMPORTANT:
#   - Placeholders use {{ DOUBLE_BRACES }}
#   - Ask user for confirmation when unsure about a value
#   - Preserve document formatting
"""

import os
import re
import json
import shutil
from datetime import datetime
from typing import Optional, List, Dict, Set
from pydantic import BaseModel, Field

try:
    from docx import Document
    from docx.shared import Inches, Pt
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class Tools:
    class Valves(BaseModel):
        templates_dir: str = Field(
            default="/data/templates",
            description="Directory containing template files"
        )
        output_dir: str = Field(
            default="/data/user-files/documents",
            description="Directory for generated documents"
        )
        backup_templates: bool = Field(
            default=True,
            description="Create backup before modifying templates"
        )
        placeholder_pattern: str = Field(
            default=r"\{\{\s*([A-Z_][A-Z0-9_]*)\s*\}\}",
            description="Regex pattern for placeholders (default: {{ PLACEHOLDER }})"
        )

    def __init__(self):
        self.citation = False
        self.valves = self.Valves()
        self._ensure_directories()

    def _ensure_directories(self):
        """Create necessary directories if they don't exist."""
        os.makedirs(self.valves.templates_dir, exist_ok=True)
        os.makedirs(self.valves.output_dir, exist_ok=True)

    def _get_template_path(self, template_name: str) -> str:
        """Get full path to a template file."""
        if not template_name.endswith('.docx'):
            template_name = f"{template_name}.docx"
        return os.path.join(self.valves.templates_dir, template_name)

    def _get_output_path(self, output_name: str) -> str:
        """Get full path for output file."""
        if not output_name.endswith('.docx'):
            output_name = f"{output_name}.docx"
        return os.path.join(self.valves.output_dir, output_name)

    def _extract_placeholders(self, text: str) -> Set[str]:
        """Extract all placeholders from text."""
        pattern = self.valves.placeholder_pattern
        return set(re.findall(pattern, text, re.IGNORECASE))

    def _replace_placeholders(self, text: str, values: Dict[str, str]) -> str:
        """Replace placeholders in text with values."""
        result = text
        for key, value in values.items():
            # Match with flexible whitespace
            pattern = r"\{\{\s*" + re.escape(key) + r"\s*\}\}"
            result = re.sub(pattern, str(value), result, flags=re.IGNORECASE)
        return result

    def list_templates(self) -> str:
        """
        List all available document templates.

        Returns:
            List of templates with their placeholders.
        """
        if not DOCX_AVAILABLE:
            return "Error: python-docx not installed. Run: pip install python-docx"

        self._ensure_directories()

        templates = []
        template_dir = self.valves.templates_dir

        if not os.path.exists(template_dir):
            return f"Templates directory not found: {template_dir}\n\nCreate templates by placing .docx files in this directory."

        for file in os.listdir(template_dir):
            if file.endswith('.docx') and not file.startswith('~'):
                file_path = os.path.join(template_dir, file)
                try:
                    doc = Document(file_path)
                    placeholders = set()

                    # Check paragraphs
                    for para in doc.paragraphs:
                        placeholders.update(self._extract_placeholders(para.text))

                    # Check tables
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                placeholders.update(self._extract_placeholders(cell.text))

                    # Check headers/footers
                    for section in doc.sections:
                        if section.header:
                            for para in section.header.paragraphs:
                                placeholders.update(self._extract_placeholders(para.text))
                        if section.footer:
                            for para in section.footer.paragraphs:
                                placeholders.update(self._extract_placeholders(para.text))

                    size = os.path.getsize(file_path)
                    size_str = f"{size / 1024:.1f} KB" if size >= 1024 else f"{size} B"

                    templates.append({
                        'name': file,
                        'size': size_str,
                        'placeholders': sorted(placeholders)
                    })

                except Exception as e:
                    templates.append({
                        'name': file,
                        'size': '?',
                        'placeholders': [],
                        'error': str(e)
                    })

        if not templates:
            return f"No templates found in: {template_dir}\n\nAdd .docx files with {{ PLACEHOLDER }} syntax."

        result = ["**Available Templates:**", "=" * 50, ""]

        for t in templates:
            result.append(f"📄 **{t['name']}** ({t['size']})")
            if t.get('error'):
                result.append(f"   ⚠️ Error: {t['error']}")
            elif t['placeholders']:
                result.append(f"   Placeholders: {', '.join(t['placeholders'])}")
            else:
                result.append("   No placeholders found")
            result.append("")

        result.append(f"Templates directory: {template_dir}")
        return "\n".join(result)

    def get_placeholders(self, template_name: str) -> str:
        """
        Get all placeholders in a template.

        Args:
            template_name: Name of the template file (with or without .docx)

        Returns:
            List of placeholders that need to be filled.
        """
        if not DOCX_AVAILABLE:
            return "Error: python-docx not installed. Run: pip install python-docx"

        template_path = self._get_template_path(template_name)

        if not os.path.exists(template_path):
            return f"Template not found: {template_path}\n\nUse list_templates() to see available templates."

        try:
            doc = Document(template_path)
            placeholders = set()
            locations = {}

            # Check paragraphs
            for i, para in enumerate(doc.paragraphs):
                found = self._extract_placeholders(para.text)
                for p in found:
                    placeholders.add(p)
                    if p not in locations:
                        locations[p] = []
                    locations[p].append(f"Paragraph {i+1}")

            # Check tables
            for t_idx, table in enumerate(doc.tables):
                for r_idx, row in enumerate(table.rows):
                    for c_idx, cell in enumerate(row.cells):
                        found = self._extract_placeholders(cell.text)
                        for p in found:
                            placeholders.add(p)
                            if p not in locations:
                                locations[p] = []
                            locations[p].append(f"Table {t_idx+1}, Row {r_idx+1}, Cell {c_idx+1}")

            # Check headers/footers
            for s_idx, section in enumerate(doc.sections):
                if section.header:
                    for para in section.header.paragraphs:
                        found = self._extract_placeholders(para.text)
                        for p in found:
                            placeholders.add(p)
                            if p not in locations:
                                locations[p] = []
                            locations[p].append("Header")
                if section.footer:
                    for para in section.footer.paragraphs:
                        found = self._extract_placeholders(para.text)
                        for p in found:
                            placeholders.add(p)
                            if p not in locations:
                                locations[p] = []
                            locations[p].append("Footer")

            if not placeholders:
                return f"No placeholders found in '{template_name}'.\n\nPlaceholders use format: {{ PLACEHOLDER_NAME }}"

            result = [
                f"**Placeholders in '{template_name}':**",
                "=" * 50,
                "",
                f"Found {len(placeholders)} placeholder(s):",
                ""
            ]

            for p in sorted(placeholders):
                locs = locations.get(p, [])
                result.append(f"  • {{ {p} }}")
                if locs:
                    result.append(f"    Located in: {', '.join(locs[:3])}")

            result.extend([
                "",
                "To fill this template, use:",
                f'fill_template("{template_name}", {{"PLACEHOLDER": "value", ...}})',
            ])

            return "\n".join(result)

        except Exception as e:
            return f"Error reading template: {str(e)}"

    def fill_template(self, template_name: str, values: str, output_name: str = "") -> str:
        """
        Fill a template with provided values and save to output.

        Args:
            template_name: Name of the template file
            values: JSON object of placeholder:value pairs, e.g.:
                   '{"COMPANY_NAME": "Acme Inc", "DATE": "2024-01-15"}'
            output_name: Name for output file (default: template_name_filled_timestamp.docx)

        Returns:
            Path to generated document.
        """
        if not DOCX_AVAILABLE:
            return "Error: python-docx not installed. Run: pip install python-docx"

        template_path = self._get_template_path(template_name)

        if not os.path.exists(template_path):
            return f"Template not found: {template_path}"

        # Parse values
        try:
            if isinstance(values, str):
                values_dict = json.loads(values)
            else:
                values_dict = values
        except json.JSONDecodeError as e:
            return f"Error parsing values JSON: {e}\n\nExpected format: {{\"PLACEHOLDER\": \"value\", ...}}"

        if not isinstance(values_dict, dict):
            return "Error: values must be a JSON object"

        # Normalize keys to uppercase
        values_dict = {k.upper(): v for k, v in values_dict.items()}

        try:
            doc = Document(template_path)
            filled_count = 0

            # Fill paragraphs
            for para in doc.paragraphs:
                if self._extract_placeholders(para.text):
                    original = para.text
                    # Process each run to preserve formatting
                    for run in para.runs:
                        if self._extract_placeholders(run.text):
                            run.text = self._replace_placeholders(run.text, values_dict)
                            filled_count += 1

            # Fill tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                if self._extract_placeholders(run.text):
                                    run.text = self._replace_placeholders(run.text, values_dict)
                                    filled_count += 1

            # Fill headers/footers
            for section in doc.sections:
                if section.header:
                    for para in section.header.paragraphs:
                        for run in para.runs:
                            if self._extract_placeholders(run.text):
                                run.text = self._replace_placeholders(run.text, values_dict)
                                filled_count += 1
                if section.footer:
                    for para in section.footer.paragraphs:
                        for run in para.runs:
                            if self._extract_placeholders(run.text):
                                run.text = self._replace_placeholders(run.text, values_dict)
                                filled_count += 1

            # Generate output filename
            if not output_name:
                base_name = template_name.replace('.docx', '')
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                output_name = f"{base_name}_filled_{timestamp}.docx"

            output_path = self._get_output_path(output_name)

            # Save document
            doc.save(output_path)

            # Check for remaining placeholders
            remaining = set()
            for para in doc.paragraphs:
                remaining.update(self._extract_placeholders(para.text))

            result = [
                f"✅ Document generated successfully!",
                "",
                f"📄 Output: {output_path}",
                f"📝 Placeholders filled: {filled_count}",
            ]

            if remaining:
                result.extend([
                    "",
                    f"⚠️ Unfilled placeholders remaining: {', '.join(remaining)}",
                    "   These were not in the provided values."
                ])

            return "\n".join(result)

        except Exception as e:
            return f"Error filling template: {str(e)}"

    def preview_template(self, template_name: str, values: str = "") -> str:
        """
        Preview what the filled template would look like (text only).

        Args:
            template_name: Name of template
            values: Optional JSON of values to preview

        Returns:
            Text preview of the document.
        """
        if not DOCX_AVAILABLE:
            return "Error: python-docx not installed"

        template_path = self._get_template_path(template_name)

        if not os.path.exists(template_path):
            return f"Template not found: {template_path}"

        try:
            doc = Document(template_path)

            # Parse values if provided
            values_dict = {}
            if values:
                try:
                    values_dict = json.loads(values) if isinstance(values, str) else values
                    values_dict = {k.upper(): v for k, v in values_dict.items()}
                except:
                    pass

            result = [
                f"**Preview of '{template_name}':**",
                "=" * 50,
                ""
            ]

            # Extract text from paragraphs
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    if values_dict:
                        text = self._replace_placeholders(text, values_dict)
                    result.append(text)

            # Extract text from tables
            for t_idx, table in enumerate(doc.tables):
                result.append(f"\n[Table {t_idx + 1}]")
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if values_dict:
                            cell_text = self._replace_placeholders(cell_text, values_dict)
                        row_text.append(cell_text)
                    result.append(" | ".join(row_text))

            if not any(r.strip() and r != "=" * 50 for r in result[2:]):
                result.append("(Document appears to be empty or contains only images)")

            return "\n".join(result)

        except Exception as e:
            return f"Error previewing template: {str(e)}"

    def create_template(self, template_name: str, content: str, placeholders: str = "") -> str:
        """
        Create a simple template from text content.

        Args:
            template_name: Name for the new template
            content: Text content with {{ PLACEHOLDERS }}
            placeholders: Optional comma-separated list of placeholder names

        Returns:
            Path to created template.
        """
        if not DOCX_AVAILABLE:
            return "Error: python-docx not installed"

        template_path = self._get_template_path(template_name)

        if os.path.exists(template_path):
            return f"Template already exists: {template_path}\n\nUse a different name or delete the existing template."

        try:
            doc = Document()

            # Add content as paragraphs
            for line in content.split('\n'):
                doc.add_paragraph(line)

            doc.save(template_path)

            # Get placeholders from content
            found = self._extract_placeholders(content)

            result = [
                f"✅ Template created: {template_path}",
                "",
            ]

            if found:
                result.append(f"Placeholders found: {', '.join(sorted(found))}")
            else:
                result.append("No placeholders found in content.")
                result.append("Add placeholders using {{ PLACEHOLDER_NAME }} syntax.")

            return "\n".join(result)

        except Exception as e:
            return f"Error creating template: {str(e)}"

    def add_template(self, source_path: str, template_name: str = "") -> str:
        """
        Copy an existing DOCX file to the templates directory.

        Args:
            source_path: Path to existing DOCX file
            template_name: Optional new name (default: use original name)

        Returns:
            Success message with placeholder info.
        """
        if not os.path.exists(source_path):
            return f"Source file not found: {source_path}"

        if not source_path.endswith('.docx'):
            return "Error: Source file must be a .docx file"

        if not template_name:
            template_name = os.path.basename(source_path)

        dest_path = self._get_template_path(template_name)

        try:
            shutil.copy2(source_path, dest_path)

            # Get placeholders
            placeholders_info = self.get_placeholders(template_name)

            return f"✅ Template added: {dest_path}\n\n{placeholders_info}"

        except Exception as e:
            return f"Error adding template: {str(e)}"

    def check_availability(self) -> str:
        """Check if the tool is properly configured."""
        status = ["Templates Tool - Status", "=" * 40, ""]

        if DOCX_AVAILABLE:
            status.append("✅ python-docx: Available")
        else:
            status.append("❌ python-docx: Not installed")
            status.append("   Run: pip install python-docx")

        status.extend([
            "",
            f"Templates directory: {self.valves.templates_dir}",
            f"Output directory: {self.valves.output_dir}",
            "",
            "Placeholder format: {{ PLACEHOLDER_NAME }}",
            "",
            "Available functions:",
            "  • list_templates() - See all templates",
            "  • get_placeholders(template) - See placeholders",
            "  • fill_template(template, values) - Generate document",
            "  • preview_template(template) - Preview content",
            "  • create_template(name, content) - Create new template",
        ])

        return "\n".join(status)
